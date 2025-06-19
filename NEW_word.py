import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx2pdf import convert
import os

# --- STEP 1: CONFIG PATHS ---

# Path to Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\aryan.singh\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

# Path to Poppler bin (for pdf2image)
poppler_path = r"C:\Users\aryan.singh\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin"  # Example: Replace with your actual Poppler path

# File paths
pdf_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new.pdf"
word_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new_output.docx"

modified_word_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new_output_modified.docx"
final_pdf_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new_output_modified.pdf"

# --- STEP 2: Convert PDF to Word via OCR ---

pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)

doc = Document()

for i, page in enumerate(pages):
    text = pytesseract.image_to_string(page)
    doc.add_paragraph(text)
    doc.add_page_break()

doc.save(word_path)
print(f"✅ Scanned PDF converted to Word: {word_path}")

# --- STEP 3: Modify the Word document ---

doc = Document(word_path)

for para in doc.paragraphs:
    if 'oldword' in para.text:
        para.text = para.text.replace('oldword', 'newword')  # 📝 Replace this with your actual words

doc.save(modified_word_path)
print(f"✍️ Modified Word file saved: {modified_word_path}")

# --- STEP 4: Convert Modified Word to PDF ---

convert(modified_word_path, final_pdf_path)
print(f"📄 Final PDF saved: {final_pdf_path}")
