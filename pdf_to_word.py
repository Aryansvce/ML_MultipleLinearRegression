import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os

# Set path to tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\aryan.singh\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

pdf_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new.pdf"
output_word_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new_output.docx"

# Path to Poppler bin folder (update this if yours is different)
poppler_path = r"C:\Users\aryan.singh\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin"  # Example: Replace with your actual Poppler path

# Convert PDF to images (one image per page)
pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)

doc = Document()

# OCR each page and write to Word
for page_num, page in enumerate(pages):
    text = pytesseract.image_to_string(page)
    doc.add_paragraph(text)
    doc.add_page_break()

# Save to Word
doc.save(output_word_path)

print(f"PDF converted to Word and saved as {output_word_path}")

# Now: Modify text in the Word file
doc = Document(output_word_path)

# Example: Replace a word (e.g., 'oldword' to 'newword')
for para in doc.paragraphs:
    if 'oldword' in para.text:
        para.text = para.text.replace('oldword', 'newword')

# Save modified document
modified_word_path = r"C:\Users\aryan.singh\Downloads\downloaded_files\patna_Delhi_new_output_modified.docx"
doc.save(modified_word_path)

print(f"Modified Word file saved as {modified_word_path}")
